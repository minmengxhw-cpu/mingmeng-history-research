#!/usr/bin/env python3
"""党政机关公文格式 docx 生成器（GB/T 9704-2012）

为本研究库今后所有对外发送的 docx 提供统一公文体例支持：
- 仿宋 GB2312 三号正文 + 28 磅固定行距 + 2 字首行缩进
- 黑体三号 / 楷体三号 / 仿宋加粗 三级标题
- 方正小标宋简体二号 公文大标题
- A4 + 上 37/下 35/左 28/右 26 mm 公文页边距
- 跨平台字体 fallback（Windows / macOS / Linux）

用法示例：
    from gov_docx import GovDoc

    doc = GovDoc()
    doc.title('民盟海外史料研究简报')
    doc.subtitle('2026 年第 5 期')
    doc.recipient('民盟上海市委各级组织：')
    doc.paragraph('为深入推进上海民盟组织建立 80 周年研究工作……')
    doc.h1('一、研究进展')
    doc.paragraph('截至 2026 年 5 月，FRUS 数据库已完成入库 416 段……')
    doc.h2('（一）核心档案')
    doc.h3('1. 民盟成立至 1944 年改名')
    doc.signature(org='民盟海外史料研究项目组', date='2026 年 5 月 16 日')
    doc.save('exports/output_公文版.docx')
"""
from __future__ import annotations
from datetime import date
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ===== 字体 fallback =====
# Word 渲染时按 east-asian / hAnsi / ascii 多字段轮询本机字体。
# 字段名按 GB/T 标准中文名书写（如 "仿宋"），ascii 给定通用拉丁字体。

FONT_FANGSONG = "仿宋"               # 正文 / 多数元数据
FONT_HEITI = "黑体"                  # 一级标题
FONT_KAITI = "楷体"                  # 二级标题
FONT_XIAOBIAOSONG = "方正小标宋简体"  # 公文大标题
FONT_ASCII = "Times New Roman"       # 西文 fallback

# 字号（pt）
SIZE_TITLE = Pt(22)         # 二号
SIZE_H1 = Pt(16)            # 三号
SIZE_BODY = Pt(16)          # 三号
SIZE_META4 = Pt(14)         # 四号（页码、抄送、印发）
SIZE_META5 = Pt(10.5)       # 五号

# 行距：28 磅固定值
LINE_SPACING_PT = 28
# 首行缩进 2 字符 ≈ 0.74 cm（三号字 16 pt × 2 ≈ 32 pt ≈ 0.83 cm）
FIRST_LINE_INDENT_CM = 0.74


def _set_run_font(run, east_asia: str = FONT_FANGSONG, ascii_font: str = FONT_ASCII,
                  size: Pt = SIZE_BODY, bold: bool = False):
    """统一设置 run 的字体（含 East-Asian fallback）"""
    run.font.name = ascii_font
    run.font.size = size
    run.font.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), east_asia)
    rFonts.set(qn('w:ascii'), ascii_font)
    rFonts.set(qn('w:hAnsi'), ascii_font)
    rFonts.set(qn('w:cs'), ascii_font)


def _set_para_line_spacing(para, line_pt: float = LINE_SPACING_PT, first_indent_cm: float = 0.0):
    """行距 28 磅固定值 + 首行缩进"""
    pPr = para.paragraph_format
    pPr.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pPr.line_spacing = Pt(line_pt)
    pPr.space_before = Pt(0)
    pPr.space_after = Pt(0)
    if first_indent_cm > 0:
        pPr.first_line_indent = Cm(first_indent_cm)


class GovDoc:
    """党政公文格式 docx 构建器（非正式行文体例：无红头无印章，
    但字体字号行距严格按 GB/T 9704-2012）"""

    def __init__(self):
        self.doc = Document()
        self._setup_page()
        self._setup_normal_style()

    def _setup_page(self):
        for section in self.doc.sections:
            section.page_width = Cm(21.0)        # A4
            section.page_height = Cm(29.7)
            section.top_margin = Mm(37)
            section.bottom_margin = Mm(35)
            section.left_margin = Mm(28)
            section.right_margin = Mm(26)

    def _setup_normal_style(self):
        st = self.doc.styles['Normal']
        st.font.name = FONT_ASCII
        st.font.size = SIZE_BODY
        rPr = st.element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.append(rFonts)
        rFonts.set(qn('w:eastAsia'), FONT_FANGSONG)
        rFonts.set(qn('w:ascii'), FONT_ASCII)

    # -------- API --------

    def title(self, text: str):
        """公文大标题：方正小标宋简体二号，居中，加粗"""
        # 标题前空 2 行
        for _ in range(2):
            p = self.doc.add_paragraph()
            _set_para_line_spacing(p)
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_para_line_spacing(p, line_pt=32)
        r = p.add_run(text)
        _set_run_font(r, east_asia=FONT_XIAOBIAOSONG, size=SIZE_TITLE, bold=True)

    def subtitle(self, text: str):
        """副标题（如「2026 年第 5 期」），三号仿宋居中"""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_para_line_spacing(p)
        r = p.add_run(text)
        _set_run_font(r, east_asia=FONT_FANGSONG, size=SIZE_H1, bold=False)
        # 副标题后空一行
        empty = self.doc.add_paragraph()
        _set_para_line_spacing(empty)

    def recipient(self, text: str):
        """主送机关：三号仿宋顶格"""
        p = self.doc.add_paragraph()
        _set_para_line_spacing(p)
        r = p.add_run(text)
        _set_run_font(r, east_asia=FONT_FANGSONG, size=SIZE_BODY, bold=False)

    def paragraph(self, text: str, indent: bool = True):
        """正文段落：三号仿宋，首行缩进 2 字符"""
        p = self.doc.add_paragraph()
        _set_para_line_spacing(p, first_indent_cm=FIRST_LINE_INDENT_CM if indent else 0.0)
        r = p.add_run(text)
        _set_run_font(r, east_asia=FONT_FANGSONG, size=SIZE_BODY, bold=False)

    def h1(self, text: str):
        """一级标题：黑体三号，左对齐，无缩进"""
        # 一级标题前空一行
        empty = self.doc.add_paragraph()
        _set_para_line_spacing(empty)
        p = self.doc.add_paragraph()
        _set_para_line_spacing(p)
        r = p.add_run(text)
        _set_run_font(r, east_asia=FONT_HEITI, size=SIZE_H1, bold=False)

    def h2(self, text: str):
        """二级标题：楷体_GB2312 三号，首行缩进 2 字符"""
        p = self.doc.add_paragraph()
        _set_para_line_spacing(p, first_indent_cm=FIRST_LINE_INDENT_CM)
        r = p.add_run(text)
        _set_run_font(r, east_asia=FONT_KAITI, size=SIZE_BODY, bold=False)

    def h3(self, text: str):
        """三级标题：仿宋_GB2312 三号加粗，首行缩进 2 字符"""
        p = self.doc.add_paragraph()
        _set_para_line_spacing(p, first_indent_cm=FIRST_LINE_INDENT_CM)
        r = p.add_run(text)
        _set_run_font(r, east_asia=FONT_FANGSONG, size=SIZE_BODY, bold=True)

    def h4(self, text: str):
        """四级标题：仿宋_GB2312 三号（不加粗），首行缩进 2 字符"""
        p = self.doc.add_paragraph()
        _set_para_line_spacing(p, first_indent_cm=FIRST_LINE_INDENT_CM)
        r = p.add_run(text)
        _set_run_font(r, east_asia=FONT_FANGSONG, size=SIZE_BODY, bold=False)

    def quote(self, text: str):
        """长引文：左右缩进，仿宋三号"""
        p = self.doc.add_paragraph()
        _set_para_line_spacing(p)
        p.paragraph_format.left_indent = Cm(1.5)
        p.paragraph_format.right_indent = Cm(1.5)
        r = p.add_run(text)
        _set_run_font(r, east_asia=FONT_FANGSONG, size=SIZE_BODY, bold=False)

    def attachment_line(self, text: str):
        """附件说明：「附件：1.xxx 2.xxx」三号仿宋"""
        p = self.doc.add_paragraph()
        _set_para_line_spacing(p, first_indent_cm=FIRST_LINE_INDENT_CM)
        r = p.add_run(text)
        _set_run_font(r, east_asia=FONT_FANGSONG, size=SIZE_BODY, bold=False)

    def signature(self, org: str, date_str: str = None):
        """发文机关署名 + 成文日期，右对齐三号仿宋"""
        if date_str is None:
            today = date.today()
            date_str = f"{today.year} 年 {today.month} 月 {today.day} 日"
        # 留 2 行空白
        for _ in range(2):
            empty = self.doc.add_paragraph()
            _set_para_line_spacing(empty)
        # 机关名
        p1 = self.doc.add_paragraph()
        p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _set_para_line_spacing(p1)
        r1 = p1.add_run(org)
        _set_run_font(r1, east_asia=FONT_FANGSONG, size=SIZE_BODY, bold=False)
        # 成文日期
        p2 = self.doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _set_para_line_spacing(p2)
        r2 = p2.add_run(date_str)
        _set_run_font(r2, east_asia=FONT_FANGSONG, size=SIZE_BODY, bold=False)

    def horizontal_line(self):
        """版记分隔线（细黑线）"""
        p = self.doc.add_paragraph()
        pPr = p._element.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), 'auto')
        pBdr.append(bottom)
        pPr.append(pBdr)

    def cc(self, text: str):
        """抄送机关：四号仿宋"""
        p = self.doc.add_paragraph()
        _set_para_line_spacing(p)
        r = p.add_run(text)
        _set_run_font(r, east_asia=FONT_FANGSONG, size=SIZE_META4, bold=False)

    def issuer(self, text: str):
        """印发机关 / 印发日期：四号仿宋"""
        p = self.doc.add_paragraph()
        _set_para_line_spacing(p)
        r = p.add_run(text)
        _set_run_font(r, east_asia=FONT_FANGSONG, size=SIZE_META4, bold=False)

    def save(self, path):
        out = Path(path)
        out.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(out)
        return out


# ============= 示例 / 自检脚本 =============
def _demo():
    """生成一份示范公文 docx，用于验证字体字号样式正确"""
    doc = GovDoc()
    doc.title('海外民盟历史文献研究')
    doc.subtitle('2026 年第 1 期 · 简报样本')

    doc.recipient('民盟上海市委各级组织：')
    doc.paragraph(
        '为深入推进上海民盟组织建立 80 周年（1946—2026）研究工作，'
        '本项目组系统收集整理了 1941 年至 1955 年期间海外一手原始档案中与中国民主同盟相关的史料。'
        '截至 2026 年 5 月，已完成 FRUS（美国对外关系文件集）与 CIA Reading Room 两大数据源的'
        '全量入库与中文翻译，共计 477 段史料，全部达到学术引用等级。现将研究进展简要汇报如下。'
    )

    doc.h1('一、研究范围与原则')
    doc.paragraph(
        '本项目严格遵循「只收录国外一手原始档案」的原则，目前已覆盖：'
        'FRUS 1941—1950 中国卷册（299 篇文档、416 段史料）；'
        'CIA Reading Room 1946—1954 民盟相关解密档案（62 篇文档、61 段史料）。'
        '所有史料均提供英文原文与人工校订的中文译文，附学术引用（BibTeX / Chicago / GB/T 7714）。'
    )

    doc.h1('二、研究平台主要功能')
    doc.h2('（一）人物索引')
    doc.paragraph(
        '按民盟史 6 个阶段整理 34 位核心人物（含上海民盟支部创始人王绍鏊、沈志远、施复亮、'
        '彭文应等），每位人物配 100—300 字简介与所有海外档案命中片段。'
    )
    doc.h2('（二）关键事件')
    doc.paragraph(
        '按民盟史阶段整理 13 个关键历史事件（1941—1949），含「上海民盟支部正式成立」'
        '「李闻血案」「黄竞武上海殉难」等节点，并支持 FRUS 与 CIA 跨数据源对比阅读。'
    )
    doc.h2('（三）月度时间线')
    doc.paragraph(
        '提供 1941—1955 年月度密度热力图，可直观查看民盟史高频时段'
        '（1946 年 7 月李闻血案、1947 年 10 月民盟被取缔、1949 年 5 月上海撤离 等）。'
    )

    doc.h1('三、与上海民盟 80 周年项目对接')
    doc.h3('1. 上海民盟支部成立（1946 年 8 月）专题')
    doc.paragraph(
        'CIA 1947 年 1 月 28 日报告间接证实「上海民盟总部」运作，1949 年起多次报告涉及上海民盟'
        '人事政策、人物档案与上海撤离事件。详见研究库 `/events/key/shanghai-branch-1946`。'
    )
    doc.h3('2. 上海民盟最知名烈士黄竞武专题')
    doc.paragraph(
        '黄竞武于 1949 年 5 月 12 日上海解放前夕殉难，CIA 同期档案'
        '《民盟成员自上海逃往香港》（1949.5.23）记录了民盟在上海的紧急转移行动。'
    )

    doc.h1('四、下一步工作')
    doc.paragraph(
        '一是推进 Wilson Center 数字档案库的民盟相关档案抓取与翻译；'
        '二是探勘 NARA（美国国家档案馆）RG 59 与 RG 226 系列；'
        '三是按需出具上海民盟 80 周年系列专题研究包。'
    )

    doc.attachment_line('附件：1. 海外档案数据源探勘报告')
    doc.attachment_line('      2. 民盟历史人物档案对照表')

    doc.signature(
        org='海外民盟历史文献研究项目组',
        date_str='2026 年 5 月 16 日',
    )

    doc.horizontal_line()
    doc.cc('抄送：相关项目负责人。')
    doc.horizontal_line()
    doc.issuer('海外民盟历史文献研究项目组印发        2026 年 5 月 16 日')

    out_path = Path(__file__).parent.parent / 'exports' / '党政公文格式示例_v1.docx'
    return doc.save(out_path)


if __name__ == '__main__':
    p = _demo()
    print(f'示范 docx 已生成: {p}')
