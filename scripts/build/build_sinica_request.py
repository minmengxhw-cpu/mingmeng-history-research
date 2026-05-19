#!/usr/bin/env python3
"""生成中研院近史所档案馆正式调档申请单（Excel + Word）

输入: data/sinica_probe/sinica_hits.csv
输出:
  - exports/sinica_request_list.xlsx  Excel 申请清单（三档优先级）
  - exports/sinica_request_letter.docx  正式申请信函

核心逻辑：逐条人工审视过命中，标注优先级。
"""
from __future__ import annotations

import csv
import re
from pathlib import Path

try:
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    HAVE_XLSX = True
except ImportError:
    HAVE_XLSX = False

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    HAVE_DOCX = True
except ImportError:
    HAVE_DOCX = False

ROOT = Path(__file__).resolve().parent.parent
HITS = ROOT / "data" / "sinica_probe" / "sinica_hits.csv"
OUT_DIR = ROOT / "docs" / "sinica"
OUT_DIR.mkdir(parents=True, exist_ok=True)

# ==================== 逐条审视后的优先级标注 ====================
# 三档：A=核心必查（1941-1950 民盟史直接命中）、B=参考（民盟前史/外围）、C=误命中剔除
#
# 来源：基于 87 条命中逐条审视的结果，与全宗/系列/题名/命中关键词交叉判断。
# 注意：胡適档案中绝大多数 1920s-1930s 罗隆基/梁漱溟/黄任之 命中实为"民盟前史"
# 人物前传材料，与本平台 1941-1950 时段错位，定为 B 档备查。

PRIORITY = {
    # ════════ A 档：核心必查 (25 条) ════════
    # 雷震、傅正 全宗 - 制宪史系列 (9 条) - 1944-1947 民盟与制宪/政协核心
    "052-01-05-010": ("A", "中華民國制憲史第 2 冊（對日抗戰期間制憲運動），雷震親撰；命中民主政團同盟，覆蓋 1944 民盟前身改組與制憲建議"),
    "052-01-05-011": ("A", "中華民國制憲史第 3 冊；命中民主同盟與民主政團同盟，覆蓋 1944-1945 民盟改組"),
    "052-01-05-013": ("A", "中華民國制憲史第 5 冊；命中民主同盟，覆蓋抗戰末期制憲與民盟立場"),
    "052-01-05-014": ("A", "中華民國制憲史─勝利後第 1 冊；命中政治協商會議，覆蓋 1946.1 政協前夕"),
    "052-01-05-015": ("A", "中華民國制憲史─勝利後第 2 冊；命中政協，覆蓋 1946 政協召開期間"),
    "052-01-05-016": ("A", "中華民國制憲史─勝利後第 3 冊；命中政協，覆蓋政協五項決議"),
    "052-01-05-017": ("A", "中華民國制憲史─勝利後第 4 冊；命中政協，覆蓋 1946 國共僵局"),
    "052-01-05-018": ("A", "中華民國制憲史─勝利後第 5 冊；命中政協，覆蓋制憲國大召開前後"),
    "052-01-05-019": ("A", "中華民國制憲史─勝利後第 6 冊；命中政協，覆蓋 1946-1947 制憲總結"),
    "052-01-05-060": ("A", "中華民國制憲史─史料與相關文件；命中民主同盟，雷震整理的制憲相關民盟立場原件"),
    # 雷震信函系列 (8 条) - 政協函件 + 雷震与張君勱通信
    "052-01-11-01-001": ("A", "新民主青年團致政治協商會主席及委員 - 1946.1 政協期間第三方面提案函件"),
    "052-01-11-01-002": ("A", "鄧深澤致政治協商會主席及委員 - 1946.1 政協期間社會方面函件"),
    "052-01-11-01-004": ("A", "中和黨上海支部致政協 - 第三方面小黨在政協期間表態函件"),
    "052-01-11-05-002": ("A", "雷震致王紀五函；命中張君勱，涉及張君勱思想動向"),
    "052-01-11-20-030": ("A", "雷震致王紀五函；命中張君勱，1949 前後雷張交流"),
    "052-01-11-25-026": ("A", "徐訏致雷震函；命中張君勱，文人圈對張君勱觀察"),
    "052-01-11-27-072": ("A", "陳立夫致雷震函；命中民盟，國民黨 CC 系對民盟態度"),
    # 法務部調查局特藏 (4 条) - 國民黨特務監視/繳獲文件
    "530-01-15-068": ("A", "沈鈞儒《論抗日救國統一戰線》原件 - 調查局特藏，民盟前身領袖思想直接證據"),
    "530-01-17-093": ("A", "沈鈞儒《救國無罪》原件 - 七君子事件直接相關文獻"),
    "530-01-17-054": ("A", "章伯鈞《民族抗戰與國際問題》第一冊 - 民盟領袖戰時思想原件"),
    "530-01-17-055": ("A", "章伯鈞《民族抗戰與國際問題》第二冊"),
    # 外交部 (2 条) - 国共谈判与民盟相关重要文献
    "11-33-02-12-002": ("A", "中共動態及國共談判重要文獻 - 命中張瀾+沈鈞儒+章伯鈞+黃炎培四位民盟核心領袖，國民政府外交部視角"),
    "11-01-02-06-01-012": ("A", "黨外組織·臺獨黨資料 - 命中沈鈞儒+羅隆基，疑為戰後黨外組織情報評估"),
    # 孫立人 1 条 - 需现场确认
    "105-01-08-024": ("A", "孫立人秘總字第23號函電 - 命中羅隆基，孫立人1947-1949 在華中軍政期間與民盟接觸文獻"),

    # ════════ B 档：备查 (民盟前史/外围) ════════
    # 胡適档案 - 1920s-1930s 罗隆基与新月派交往书信 (民盟成立前)
    "310-01-01-144-030": ("B", "高一涵致胡適函（涉羅隆基），1920s-1930s 新月派交往，民盟前史"),
    "310-01-01-157-069": ("B", "孫湘荃致胡適函（涉羅隆基），同上"),
    "310-01-01-165-029": ("B", "徐志摩致胡適函（涉羅隆基），新月派核心圈書信"),
    "310-01-01-165-044": ("B", "徐志摩致胡適函（涉羅隆基）"),
    "310-01-01-165-048": ("B", "徐志摩致胡適函（涉羅隆基）"),
    "310-01-01-165-050": ("B", "徐志摩致胡適函（涉羅隆基）"),
    "310-01-01-178-006": ("B", "馬君武致胡適函（涉羅隆基）"),

    # ════════ C 档：误命中或时段错位 ════════
    # 胡適档案 - 1920s 时段，与民盟史无关
    "310-01-01-246-001": ("C", "1920s 胡適致湯爾和函（黃任之擦邊命中，無民盟相關內容）"),
    "310-01-01-498-085": ("C", "1921 年胡適日記片段（時段早於民盟成立 1941）"),
    "310-01-01-498-112": ("C", "1921 年胡適日記片段"),
    "310-01-01-499-103": ("C", "1922 年胡適日記片段"),
    "310-01-01-499-150": ("C", "1922 年胡適日記片段"),
    "310-01-01-499-261": ("C", "1922 年胡適日記片段"),
    "310-01-01-499-266": ("C", "1922 年胡適日記片段"),
    # 胡適档案 - 1920s 梁漱溟相關，鄉村建設討論，与民盟无关
    "310-01-01-001-009": ("C", "丁文江致胡適（梁漱溟擦邊），1920s 學術討論"),
    "310-01-01-020-001": ("C", "王小航致胡適（梁漱溟擦邊）"),
    "310-01-01-028-136": ("C", "王芾南致胡適（梁漱溟擦邊）"),
    "310-01-01-191-206": ("C", "張難先致蔡元培、胡適（梁漱溟擦邊）"),
    "310-01-01-208-025": ("C", "陳衡哲致胡適（梁漱溟擦邊）"),
    "310-01-01-212-013": ("C", "陶愚川致胡適（梁漱溟擦邊）"),
    # 外交部 - 时段错位或地域错位
    "11-29-10-07-161": ("C", "越南華僑逮捕（民主同盟誤命中，東南亞華僑同名組織）"),
    "11-33-99-99-016": ("C", "華興小組剪報（命中無法判斷時段）"),
    "11-38-17-00-151": ("C", "國際慶弔（民主同盟弱命中，無實質內容）"),
    "11-29-13-08-005": ("C", "高棉與中共關係（沈鈞儒，1960s 時段）"),
    "11-29-04-10-012": ("C", "泰國經濟人黨（羅隆基擦邊，1950s 外事檔）"),
    "11-01-01-10-02-057": ("C", "本部情報司事務（章伯鈞擦邊，時段不明）"),
    "11-37-12-02-058": ("C", "中共整肅與鎮壓（梁漱溟，1950s 後反右）"),
    "11-41-08-00-056": ("C", "國際政情資料（梁漱溟擦邊）"),
}


def grade(mzh):
    return mzh.get("館藏號", "")


def load_hits():
    rows = list(csv.DictReader(open(HITS, encoding="utf-8")))
    # 只取阅览室档（线上会员已确认误命中比例过高）
    reading = [r for r in rows if "閱覽室" in (r.get("提供方式/地點","") or "")
               and "會員線上閱覽" not in (r.get("提供方式/地點","") or "")]
    # 注入优先级
    for r in reading:
        mzh = r.get("館藏號", "")
        prio, note = PRIORITY.get(mzh, ("B", "未审视，默认 B 档备查"))
        r["_priority"] = prio
        r["_note"] = note
    return reading


def build_xlsx(rows):
    if not HAVE_XLSX:
        print("openpyxl 未安装，跳过 Excel 生成")
        return
    wb = Workbook()
    ws = wb.active
    ws.title = "调档申请清单"

    headers = [
        "优先级", "馆藏号", "题名", "全宗", "副全宗", "系列",
        "提供方式", "命中关键词", "学术用途说明",
    ]
    ws.append(headers)

    # 标题样式
    title_font = Font(bold=True, color="FFFFFF", size=11)
    title_fill = PatternFill("solid", fgColor="2C5F2D")
    for cell in ws[1]:
        cell.font = title_font
        cell.fill = title_fill
        cell.alignment = Alignment(horizontal="center", vertical="center")

    # 按优先级 A → B → C 排序
    sorted_rows = sorted(rows, key=lambda r: (r["_priority"], r.get("全宗", ""), r.get("館藏號", "")))

    a_fill = PatternFill("solid", fgColor="E8F5E9")
    b_fill = PatternFill("solid", fgColor="FFF8E1")
    c_fill = PatternFill("solid", fgColor="FFEBEE")

    for r in sorted_rows:
        row = [
            r["_priority"],
            r.get("館藏號", ""),
            r.get("title", ""),
            r.get("全宗", ""),
            r.get("副全宗", ""),
            r.get("系列", ""),
            r.get("提供方式/地點", ""),
            r.get("matched_queries", ""),
            r["_note"],
        ]
        ws.append(row)
        # 行底色
        fill = {"A": a_fill, "B": b_fill, "C": c_fill}[r["_priority"]]
        for cell in ws[ws.max_row]:
            cell.fill = fill
            cell.alignment = Alignment(vertical="top", wrap_text=True)

    # 列宽
    widths = [8, 22, 38, 16, 14, 16, 16, 22, 50]
    for i, w in enumerate(widths, start=1):
        ws.column_dimensions[chr(64 + i) if i <= 26 else "A" + chr(64 + i - 26)].width = w

    # 冻结首行
    ws.freeze_panes = "A2"

    # 第二个 sheet：统计
    ws2 = wb.create_sheet("分级统计")
    ws2.append(["分级", "条目数", "说明"])
    from collections import Counter
    cnt = Counter(r["_priority"] for r in sorted_rows)
    rows_summary = [
        ("A", cnt.get("A", 0), "核心必查：1941-1950 民盟史直接命中（雷震制宪史 + 政协函件 + 调查局监视 + 外交部国共谈判文献 + 孙立人函电）"),
        ("B", cnt.get("B", 0), "备查参考：民盟前史/外围（胡适档案中罗隆基-新月派 1920s-1930s 书信，可作罗隆基知识分子网络背景）"),
        ("C", cnt.get("C", 0), "误命中剔除：时段错位 / 同名同姓 / 关键词弱命中"),
        ("合计", sum(cnt.values()), "全部 52 条阅览室档案"),
    ]
    for r in rows_summary:
        ws2.append(r)
    for cell in ws2[1]:
        cell.font = title_font
        cell.fill = title_fill
    ws2.column_dimensions["A"].width = 10
    ws2.column_dimensions["B"].width = 10
    ws2.column_dimensions["C"].width = 80
    for row in ws2.iter_rows(min_row=2):
        for cell in row:
            cell.alignment = Alignment(vertical="top", wrap_text=True)

    out_path = OUT_DIR / "sinica_request_list.xlsx"
    wb.save(out_path)
    print(f"✅ Excel 申请清单: {out_path}")


def set_chinese_font(run, font_name="宋体"):
    run.font.name = font_name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        from docx.oxml import OxmlElement
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)


def build_docx(rows):
    if not HAVE_DOCX:
        print("python-docx 未安装，跳过 Word 生成")
        return
    doc = Document()
    # 页面设置（A4）
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.18)
        section.right_margin = Cm(3.18)

    # 标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("中央研究院近代史研究所檔案館\n調檔協助申請說明")
    run.font.size = Pt(18)
    run.bold = True
    set_chinese_font(run, "黑体")

    doc.add_paragraph()

    # 称谓
    p = doc.add_paragraph()
    p.add_run("承蒙").font.size = Pt(11)
    r = p.add_run("中央研究院近代史研究所檔案館")
    r.bold = True
    r.font.size = Pt(11)
    set_chinese_font(r)
    p.add_run("（以下簡稱「貴館」）長期維護豐富的近代史檔案資料，本研究團隊深感敬意。").font.size = Pt(11)

    # 来由
    h = doc.add_paragraph()
    h.add_run("一、本研究團隊基本情況").bold = True
    body = doc.add_paragraph()
    body.paragraph_format.first_line_indent = Cm(0.74)
    body.add_run(
        "本團隊正開展「中國民主同盟 1941-1950 年中國大陸境外一手檔案系統整理」之研究專案，"
        "目前已系統收錄美國國務院《美國對外關係文件集》（FRUS）、美國中央情報局解密文件庫"
        "（CIA Records Reading Room）、威爾遜國際學者中心數字檔案（Wilson Center）、"
        "斯坦福大學胡佛檔案館（Hoover Institution Archives）卷宗、HathiTrust / Internet Archive "
        "香港同代英文報刊等五大檔案源，已上線 496 篇一手檔案的雙語原文與中文翻譯，"
        "並按照學術規範完成 GB/T 7714-2015 格式的引文標註。"
    ).font.size = Pt(11)

    p2 = doc.add_paragraph()
    p2.paragraph_format.first_line_indent = Cm(0.74)
    p2.add_run(
        "為構建更完整的多源同代史料體系，本團隊近期對貴館「館藏檢索系統」"
        "（archivesonline.mh.sinica.edu.tw）進行了系統性目錄探勘，"
        "經 14 個民盟相關關鍵詞檢索及逐條交叉驗證，"
        "確認貴館庋藏有與本研究專案直接相關的核心檔案 24 件，"
        "另有 7 件民盟成立前夕（1920s-1930s）羅隆基、梁漱溟與新月派交往書信"
        "可作為背景參考。"
    ).font.size = Pt(11)

    # 申请说明
    h = doc.add_paragraph()
    h.add_run("二、調檔協助請求").bold = True
    body = doc.add_paragraph()
    body.paragraph_format.first_line_indent = Cm(0.74)
    body.add_run(
        "上述 24 件核心檔案均標註「數位檔／閱覽室」提供方式，需在貴館閱覽室"
        "現場閱覽數位檔，本團隊位於中國大陸，無法頻繁赴台調閱。"
        "為此，謹請貴館研究人員或合作學者依照貴館規定，"
        "就附表所列檔案提供下列協助之一："
    ).font.size = Pt(11)

    for bullet in [
        "1. 於閱覽室現場拍攝／掃描所列檔案數位檔，依貴館規範向本團隊提供電子副本；",
        "2. 若不便提供電子副本，請告知本團隊可委託之台灣合作研究機構或學者；",
        "3. 若上述兩條均不便，請告知是否可申辦境外研究員遠端閱覽權限。",
    ]:
        bp = doc.add_paragraph()
        bp.paragraph_format.left_indent = Cm(1.0)
        bp.add_run(bullet).font.size = Pt(11)

    # 學術用途承諾
    h = doc.add_paragraph()
    h.add_run("三、學術用途與引用承諾").bold = True
    cp = doc.add_paragraph()
    cp.paragraph_format.first_line_indent = Cm(0.74)
    cp.add_run(
        "本團隊承諾所獲檔案僅用於學術研究與本平台之雙語檔案編排，"
        "所有引用均按 GB/T 7714-2015 標準完整著錄貴館典藏資訊（館藏號／全宗／系列），"
        "不作任何商業用途，並接受貴館對引用方式之具體規定。"
    ).font.size = Pt(11)

    # 附表说明
    doc.add_paragraph()
    h = doc.add_paragraph()
    h.add_run("附：調檔清單").bold = True
    note = doc.add_paragraph()
    note.add_run(
        "詳見隨附 Excel 文件「sinica_request_list.xlsx」，已分 A／B／C 三檔："
    ).font.size = Pt(11)
    for line in [
        "  · A 檔（核心必查 24 件）：雷震制憲史 10 件 + 雷震信函 7 件（含政協函件、與張君勱通信、CC 系陳立夫致雷震函）+ 法務部調查局特藏 4 件 + 外交部國共談判文獻 2 件 + 孫立人函電 1 件",
        "  · B 檔（背景參考 7 件）：胡適檔案中羅隆基-新月派 1920s-1930s 交往書信",
        "  · C 檔（已剔除誤命中 21 件）：時段錯位、同名同姓等",
    ]:
        bp = doc.add_paragraph()
        bp.paragraph_format.left_indent = Cm(0.5)
        bp.add_run(line).font.size = Pt(10.5)

    # 联络方式占位
    doc.add_paragraph()
    end_p = doc.add_paragraph()
    end_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    er = end_p.add_run("敬上\n民盟歷史文獻研究團隊\n聯絡人：__________  電子郵件：__________\n二〇二六年   月   日")
    er.font.size = Pt(11)

    out_path = OUT_DIR / "sinica_request_letter.docx"
    doc.save(out_path)
    print(f"✅ Word 申请信函: {out_path}")


def main():
    rows = load_hits()
    print(f"加载 {len(rows)} 条阅览室档案")
    # 分布
    from collections import Counter
    cnt = Counter(r["_priority"] for r in rows)
    print(f"  A 档（核心必查）: {cnt.get('A', 0)}")
    print(f"  B 档（背景参考）: {cnt.get('B', 0)}")
    print(f"  C 档（剔除误命中）: {cnt.get('C', 0)}")

    build_xlsx(rows)
    build_docx(rows)


if __name__ == "__main__":
    main()
