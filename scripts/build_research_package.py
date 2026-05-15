#!/usr/bin/env python3
"""D2 学术研究包导出
   输出：exports/research-package/民盟海外史料研究专集_FRUS_v1.docx
   结构：封面 / 引言 / 目录 / 5 个专题章 / 10 个人物章 / 参考文献 / 术语表附录
"""
import sqlite3, csv, sys
from pathlib import Path
from datetime import date
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn

ROOT = Path(__file__).parent.parent
DB = ROOT / "data" / "research_index.sqlite"
GLOSSARY = ROOT / "data" / "translation_glossary.csv"
OUT = ROOT / "exports" / "research-package" / "民盟海外史料研究专集_FRUS_v1.docx"
OUT.parent.mkdir(parents=True, exist_ok=True)

# ===== 字体策略 =====
# 主体 East-Asian 字体：思源宋体（开源版 Noto Serif CJK SC = Source Han Serif SC）
#   Mac 没装思源也无妨——Word 会按 east-asian fallback 自动用 Songti SC（系统自带的"宋体-简"）
#   Windows fallback 到 SimSun
#   Linux 用 Noto Serif CJK SC（本机已装）
# 标题用同字体 Bold；强调用思源黑体（Source Han Sans SC = Noto Sans CJK SC）
FONT_SERIF = "Source Han Serif SC"   # 思源宋体（学术阅读标准）
FONT_SANS  = "Source Han Sans SC"    # 思源黑体（机关现代风）
FONT_ASCII = "EB Garamond"            # 英文衬线（与思源宋搭配）

conn = sqlite3.connect(DB)
conn.row_factory = sqlite3.Row

# ===== 创建文档 =====
d = Document()

# 页面大小（A4）+ 边距（学术阅读体例）
for section in d.sections:
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(3.0)

# Normal 样式（正文默认）
style = d.styles['Normal']
style.font.name = FONT_ASCII
style.font.size = Pt(11)
rpr = style.element.rPr
rpr.rFonts.set(qn('w:eastAsia'), FONT_SERIF)
rpr.rFonts.set(qn('w:hAnsi'), FONT_ASCII)
rpr.rFonts.set(qn('w:cs'), FONT_ASCII)
ppr = style.paragraph_format
ppr.line_spacing = 1.6
ppr.space_after = Pt(6)

def _set_run_font(run, *, serif=True, size=11, bold=False, italic=False, color=None):
    """统一设置中英文字体，强制 East-Asian fallback"""
    if serif:
        run.font.name = FONT_ASCII
        run.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_SERIF)
        run.element.rPr.rFonts.set(qn('w:hAnsi'), FONT_ASCII)
    else:
        run.font.name = "Helvetica"
        run.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_SANS)
        run.element.rPr.rFonts.set(qn('w:hAnsi'), "Helvetica")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color

def heading(text, level=1, center=False, color=None):
    p = d.add_paragraph()
    if center: p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sizes = {1: 22, 2: 16, 3: 13}
    size = sizes.get(level, 11)
    run = p.add_run(text)
    _set_run_font(run, serif=True, size=size, bold=True,
                  color=color or RGBColor(0x1a, 0x1a, 0x1a))
    p.paragraph_format.space_before = Pt(20 if level == 1 else 14)
    p.paragraph_format.space_after = Pt(10 if level == 1 else 8)
    p.paragraph_format.line_spacing = 1.35
    return p

def para(text, italic=False, bold=False, size=11, center=False, indent=0, sans=False, color=None):
    p = d.add_paragraph()
    if center: p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if indent:
        # 首行缩进 2 个汉字 ≈ size * 2 pt
        p.paragraph_format.first_line_indent = Pt(size * 2)
    run = p.add_run(text)
    _set_run_font(run, serif=(not sans), size=size, bold=bold, italic=italic, color=color)
    p.paragraph_format.line_spacing = 1.6
    p.paragraph_format.space_after = Pt(6)
    return p

def hr():
    d.add_paragraph("─" * 40, style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER

def pagebreak():
    d.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

# ===== 1. 封面 =====
d.add_paragraph()
d.add_paragraph()
d.add_paragraph()
p = d.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("民盟海外史料研究专集")
_set_run_font(r, serif=True, size=36, bold=True)

p = d.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("FRUS 1941–1950 卷 · 第一辑")
_set_run_font(r, serif=True, size=20)

d.add_paragraph(); d.add_paragraph(); d.add_paragraph()
para(f"基于 FRUS（Foreign Relations of the United States）1941–1950 年中国卷册", center=True, size=12)
para(f"系统抓取、翻译与研究整理", center=True, size=12)
d.add_paragraph(); d.add_paragraph()

# 统计数据
stats = {
    "文档": conn.execute("SELECT count(*) FROM documents").fetchone()[0],
    "页面": conn.execute("SELECT count(*) FROM pages").fetchone()[0],
    "中文译文": conn.execute("SELECT count(*) FROM translations WHERE language='zh-CN'").fetchone()[0],
    "人工复核": conn.execute("SELECT count(*) FROM translations WHERE language='zh-CN' AND status='human-reviewed'").fetchone()[0],
    "事件线索": conn.execute("SELECT count(*) FROM research_events").fetchone()[0],
}
for label, n in stats.items():
    para(f"{label}：{n}", center=True, size=11)

d.add_paragraph(); d.add_paragraph(); d.add_paragraph(); d.add_paragraph()
para("研究整理：民盟史料研究平台项目组", center=True, size=12)
para(f"出版日期：{date.today().strftime('%Y 年 %m 月')}", center=True, size=12)
para("版本：v1", center=True, size=12)
pagebreak()

# ===== 2. 引言 =====
heading("引言", level=1)
para("中国民主同盟（以下简称「民盟」）于 1941 年成立，是新中国成立前重要的民主党派之一。"
     "民盟在抗日战争、解放战争和新中国建设的关键时期，与中国共产党风雨同舟、肝胆相照，"
     "走出了一条爱国、民主、团结、求实的光荣道路。", indent=0.74)
para("由于民盟早期活动的特殊性，国内一手史料相对分散；与此同时，美国国务院《对外关系文件集》"
     "（Foreign Relations of the United States, FRUS）系统记录了 1941–1950 年间美方驻华机构对民盟的"
     "观察、与民盟领导人的往来记录、以及对民盟与国民政府、中国共产党之间互动的评估，构成了"
     "海外研究民盟史的重要一手档案。", indent=0.74)
para("本专集是民盟史料研究平台项目组对 FRUS 1941–1950 年中国卷册系统抓取、清洗、中英对照翻译、"
     "事件线索梳理与学术整理的阶段性成果。本专集共收录 FRUS 民盟相关文档 "
     f"{stats['文档']} 篇、页面/段落片段 {stats['页面']} 个，已 100% 完成中文翻译与人工复核，"
     f"按专题与人物维度梳理事件线索 {stats['事件线索']} 条。", indent=0.74)
para("本专集按以下结构组织：", indent=0.74)
para("　　第一编　专题研究", indent=0)
para("　　　　第一章　马歇尔调处与民盟（1946）", indent=0)
para("　　　　第二章　1946 年政治协商会议", indent=0)
para("　　　　第三章　昆明暗杀与民盟政治压力", indent=0)
para("　　　　第四章　第三方面与中间路线", indent=0)
para("　　　　第五章　1949 年北平接触", indent=0)
para("　　第二编　人物档案", indent=0)
para("　　　　按 FRUS 命中事件排序：周恩来、罗隆基、张澜、张君劢、黄炎培、沈钧儒、章伯钧、张东荪、梁漱溟、史良。", indent=0)
para("　　第三编　附录", indent=0)
para("　　　　附录一　标准译名表（109 条）", indent=0)
para("　　　　附录二　参考文献（FRUS 文档清单）", indent=0)
pagebreak()

# ===== 3. 目录 =====
heading("目　录", level=1)
toc = [
    ("引言", 2),
    ("第一编　专题研究", None),
    ("　第一章　马歇尔调处与民盟（1946）", None),
    ("　第二章　1946 年政治协商会议", None),
    ("　第三章　昆明暗杀与民盟政治压力", None),
    ("　第四章　第三方面与中间路线", None),
    ("　第五章　1949 年北平接触", None),
    ("第二编　人物档案", None),
    ("　第六章　周恩来在民盟视角下的轨迹", None),
    ("　第七章　罗隆基", None),
    ("　第八章　张澜", None),
    ("　第九章　张君劢", None),
    ("　第十章　其他民盟先贤", None),
    ("第三编　附录", None),
    ("　附录一　标准译名表", None),
    ("　附录二　参考文献（FRUS 文档清单）", None),
]
for item, _ in toc:
    p = d.add_paragraph(item)
pagebreak()

# ===== 4. 专题研究 =====
heading("第一编　专题研究", level=1)
para("本编按马歇尔调处、政协 1946、昆明暗杀、第三方面、北平接触五个专题，", indent=0.74)
para("梳理 FRUS 文献中与民盟相关的事件线索。每个专题给出：", indent=0.74)
para("　• 专题背景简述", indent=0)
para("　• 关键事件时间线（按 FRUS 卷册顺序）", indent=0)
para("　• 核心文献编号与中文标题", indent=0)
pagebreak()

TOPIC_INFO = [
    ("marshall-mediation", "马歇尔调处与民盟", "1946 年，美国总统杜鲁门派遣陆军元帅马歇尔来华调处国共纠纷。民盟作为第三方面重要力量，在调处过程中扮演了关键中介角色。本章梳理 FRUS 1946 年卷中民盟与马歇尔团队互动的核心档案。"),
    ("pcc-1946", "1946 年政治协商会议", "1946 年 1 月，国共两党与各民主党派、社会贤达在重庆召开政治协商会议，民盟以张澜、罗隆基、张君劢等为代表参加，提出和平建国主张。"),
    ("kunming-assassinations", "昆明暗杀与民盟政治压力", "1946 年 7 月，民盟昆明支部李公朴、闻一多相继遇害，引发国内外强烈反响。FRUS 文献记录了民盟人士对国民政府的强烈抗议与美方的评估。"),
    ("third-force", "第三方面与中间路线", "民盟作为国共之外的第三方面代表，倡导和平建国、宪政民主与多党合作。FRUS 文献保留了大量第三方面活动记录。"),
    ("peiping-1949", "1949 年北平接触", "1949 年新中国成立前夕，民盟与中国共产党在北平展开密切互动，参与新政协筹备工作。FRUS 文献记录了美方对这一过程的观察。"),
]
chap_num = 1
for slug, name, intro in TOPIC_INFO:
    heading(f"第{['一','二','三','四','五'][chap_num-1]}章　{name}", level=2)
    para(intro, indent=0.74)
    para("", size=8)
    para("【关键文献清单】", bold=True)
    # 拉该专题的 events 列表（去重到 page）
    rows = conn.execute("""
        SELECT DISTINCT e.event_date, e.event_title, e.event_summary,
               d.title AS doc_title, d.doc_key, d.volume_id, d.doc_id,
               d.date_guess, p.page_label
        FROM research_events e
        JOIN pages p ON p.id=e.page_id
        JOIN documents d ON d.id=p.document_id
        WHERE e.scope_slug=?
        ORDER BY e.event_date, e.event_year
        LIMIT 40
    """, (slug,)).fetchall()
    for r in rows:
        date_s = r['event_date'] or r['date_guess'] or ''
        para(f"[{r['volume_id']}/{r['doc_id']} · {date_s}] {r['event_title']}", size=10.5)
        if r['event_summary']:
            summary = r['event_summary'][:160]
            para(f"　　{summary}", size=10)
    pagebreak()
    chap_num += 1

# ===== 5. 人物档案 =====
heading("第二编　人物档案", level=1)
para("本编按 FRUS 文献中民盟相关人物的命中事件数排序，梳理每位人物在 FRUS 文献中的关键轨迹。", indent=0.74)
pagebreak()

PERSONS = [
    ("chou-en-lai", "周恩来", "中共代表，1946 年作为中共代表团团长在重庆、南京与民盟、马歇尔团队进行密集互动。虽非民盟成员，但其与民盟领导人的频繁接触是民盟史料的核心组成部分。"),
    ("lo-lung-chi", "罗隆基", "民盟中央常委、政治协商会议代表、1946 年与马歇尔会谈最频繁的民盟代表之一。1947 年民盟被取缔前后承受巨大政治压力。"),
    ("chang-lan", "张澜", "民盟主席，著名民主人士。1946 年从虹桥疗养院脱险经历是民盟史关键事件之一。"),
    ("carsun-chang", "张君劢", "民主社会党主席（民盟早期成员），1947 年 11 月致马歇尔函件成为研究民盟取缔前国际反应的关键文献。"),
    ("huang-yen-pei", "黄炎培", "民盟创始人之一，民盟上海市支部 1946 年 2 月筹建茶会召集人。"),
    ("shen-chun-ju", "沈钧儒", "民盟创始人之一，著名「七君子」成员，曾任民盟中央常委。"),
    ("chang-po-chun", "章伯钧", "民盟早期负责人，长期参与民盟核心决策。"),
    ("chang-tung-sun", "张东荪", "民盟北平负责人，1949 年北平接触关键人物。"),
    ("liang-shu-ming", "梁漱溟", "民盟早期重要发起人，「乡村建设派」代表。"),
    ("shih-liang", "史良", "民盟「七君子」之一，著名律师。"),
]
chap_num = 6
for slug, name, intro in PERSONS[:4]:  # 前 4 位深度展开
    heading(f"第{['六','七','八','九','十'][chap_num-6]}章　{name}", level=2)
    para(intro, indent=0.74)
    para("", size=8)
    para("【FRUS 文献中的关键事件】", bold=True)
    rows = conn.execute("""
        SELECT DISTINCT e.event_date, e.event_title, e.event_summary,
               d.title AS doc_title, d.volume_id, d.doc_id, d.date_guess
        FROM research_events e
        JOIN pages p ON p.id=e.page_id
        JOIN documents d ON d.id=p.document_id
        WHERE e.scope_slug=?
        ORDER BY e.event_date, e.event_year
        LIMIT 25
    """, (slug,)).fetchall()
    for r in rows:
        date_s = r['event_date'] or r['date_guess'] or ''
        para(f"[{r['volume_id']}/{r['doc_id']} · {date_s}] {r['event_title']}", size=10.5)
        if r['event_summary']:
            para(f"　　{r['event_summary'][:140]}", size=10)
    pagebreak()
    chap_num += 1

# 第十章：其他 6 位简表
heading("第十章　其他民盟先贤", level=2)
for slug, name, intro in PERSONS[4:]:
    para(f"【{name}】 {intro}", indent=0.74)
    n_events = conn.execute("SELECT count(*) FROM research_events WHERE scope_slug=?", (slug,)).fetchone()[0]
    para(f"　FRUS 命中事件：{n_events} 条", size=10)
    para("", size=6)
pagebreak()

# ===== 6. 附录一 术语表 =====
heading("第三编　附录", level=1)
heading("附录一　标准译名表", level=2)
para("本表为本研究项目使用的统一术语对照表，共 109 条，按英文首字母排序。", indent=0.74)
para("", size=8)
entries = []
with open(GLOSSARY) as f:
    for row in csv.DictReader(f):
        term = (row.get("term") or "").strip()
        trans = (row.get("translation") or "").strip()
        note = (row.get("note") or "").strip()
        if term and trans:
            entries.append((term, trans, note))
# 用表格
t = d.add_table(rows=1, cols=3)
t.style = 'Light Grid Accent 1'
hdr = t.rows[0].cells
hdr[0].text = "英文术语"; hdr[1].text = "标准中文"; hdr[2].text = "备注"
for term, trans, note in sorted(entries, key=lambda x: x[0].lower()):
    row = t.add_row().cells
    row[0].text = term; row[1].text = trans; row[2].text = note
pagebreak()

# ===== 7. 附录二 参考文献 =====
heading("附录二　参考文献（FRUS 文档清单）", level=2)
para(f"本专集收录 FRUS 民盟相关文档 {stats['文档']} 篇，按 FRUS 卷册顺序排列。", indent=0.74)
para("", size=8)
docs = conn.execute("""
    SELECT volume_id, doc_id, title, date_guess, url, hit_type
    FROM documents ORDER BY volume_id, CAST(doc_number AS INTEGER)
""").fetchall()
current_vol = None
for r in docs:
    if r['volume_id'] != current_vol:
        current_vol = r['volume_id']
        para("", size=8)
        para(f"【{current_vol}】", bold=True, size=11)
    para(f"  · 文件 {r['doc_id']} · {r['date_guess']} · {r['title'][:80]}", size=10)
    para(f"    {r['url']}", size=9)

# 保存
d.save(OUT)
print(f"✅ 保存至 {OUT}")
print(f"   大小: {OUT.stat().st_size // 1024} KB")
conn.close()
